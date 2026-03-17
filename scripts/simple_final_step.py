#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
商务专家skill最后一步：将现有Markdown报告转换为Word文档
"""

import os
import sys
from pathlib import Path

def main():
    print("=" * 60)
    print("商务专家skill - 最后一步：生成Word文档")
    print("=" * 60)
    
    # 输入文件
    md_path = r"C:\Users\mingh\.openclaw\workspace\skills\business-expert\scripts\reports_fixed\CBD_经营分析报告_20260317.md"
    docx_path = r"C:\Users\mingh\.openclaw\workspace\skills\business-expert\scripts\reports_fixed\CBD_经营分析报告_20260317_final.docx"
    
    print(f"输入文件: {md_path}")
    print(f"输出文件: {docx_path}")
    
    # 检查输入文件
    if not os.path.exists(md_path):
        print(f"❌ Markdown文件不存在: {md_path}")
        return 1
    
    # 读取Markdown内容
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        print(f"✅ 成功读取Markdown文件")
        print(f"   文件大小: {len(markdown_content):,} 字符")
    except Exception as e:
        print(f"❌ 读取Markdown文件失败: {e}")
        return 1
    
    # 导入markdown-to-word-skill
    try:
        skill_path = Path(r"C:\Users\mingh\.openclaw\workspace\skills\markdown-to-word-skill\scripts")
        sys.path.insert(0, str(skill_path))
        
        from md2docx import MarkdownToDocxConverter
        
        print("✅ 成功导入markdown-to-word-skill")
        
        # 使用商业模板
        template_dir = Path(r"C:\Users\mingh\.openclaw\workspace\skills\markdown-to-word-skill\templates")
        template_path = template_dir / "business.docx"
        
        if template_path.exists():
            print(f"✅ 使用商业模板: {template_path}")
            converter = MarkdownToDocxConverter(template_path=str(template_path), debug=True)
        else:
            print("⚠️  商业模板不存在，使用默认模板")
            converter = MarkdownToDocxConverter(debug=True)
        
        # 转换文档
        print("🔄 正在转换Markdown为Word...")
        converter.convert(markdown_content)
        converter.doc.save(docx_path)
        
        # 检查结果
        if os.path.exists(docx_path):
            file_size = os.path.getsize(docx_path)
            print(f"\n" + "=" * 60)
            print("✅ 转换成功!")
            print("=" * 60)
            print(f"Word文档: {docx_path}")
            print(f"文件大小: {file_size:,} 字节")
            print(f"生成时间: 2026-03-17 08:39")
            print("=" * 60)
            
            # 显示文档预览
            print("\n📄 文档内容预览:")
            print("-" * 40)
            lines = markdown_content.split('\n')[:15]
            for line in lines[:15]:
                if line.strip():
                    print(f"  {line[:80]}{'...' if len(line) > 80 else ''}")
            print("-" * 40)
            
            return 0
        else:
            print("❌ Word文档生成失败")
            return 1
            
    except ImportError as e:
        print(f"❌ 导入markdown-to-word-skill失败: {e}")
        print("尝试使用python-docx直接转换...")
        
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # 简单转换
            lines = markdown_content.split('\n')
            for line in lines:
                if line.strip():
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=0)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=1)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=2)
                    else:
                        doc.add_paragraph(line)
            
            doc.save(docx_path)
            
            if os.path.exists(docx_path):
                file_size = os.path.getsize(docx_path)
                print(f"✅ 使用python-docx转换成功: {docx_path}")
                print(f"   文件大小: {file_size:,} 字节")
                return 0
            else:
                print("❌ python-docx转换失败")
                return 1
                
        except Exception as e2:
            print(f"❌ python-docx转换也失败: {e2}")
            return 1
    
    except Exception as e:
        print(f"❌ 转换过程失败: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(main())