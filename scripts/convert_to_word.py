#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown报告转换为Word文档
"""

import os
import sys
from pathlib import Path

# 添加markdown-to-word-skill的路径
skill_path = Path(r"C:\Users\mingh\.openclaw\workspace\skills\markdown-to-word-skill\scripts")
sys.path.insert(0, str(skill_path))

try:
    from md2docx import MarkdownToDocxConverter
    print("成功导入MarkdownToDocxConverter")
except ImportError as e:
    print(f"导入失败: {e}")
    print("尝试直接使用python-docx转换...")
    
    # 如果导入失败，使用简单的转换方法
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    class SimpleMarkdownToDocxConverter:
        """简单的Markdown转Word转换器"""
        
        def __init__(self):
            self.doc = Document()
            
        def convert(self, markdown_text: str, output_path: str):
            """简单转换Markdown为Word"""
            lines = markdown_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 处理标题
                if line.startswith('# '):
                    self.doc.add_heading(line[2:], level=0)
                elif line.startswith('## '):
                    self.doc.add_heading(line[3:], level=1)
                elif line.startswith('### '):
                    self.doc.add_heading(line[4:], level=2)
                elif line.startswith('#### '):
                    self.doc.add_heading(line[5:], level=3)
                # 处理表格行
                elif '|' in line and '---' not in line:
                    # 简单表格处理
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if cells:
                        p = self.doc.add_paragraph()
                        for cell in cells:
                            p.add_run(cell).bold = True
                            if cell != cells[-1]:
                                p.add_run(' | ')
                else:
                    # 普通段落
                    self.doc.add_paragraph(line)
            
            # 保存文档
            self.doc.save(output_path)
            print(f"文档已保存: {output_path}")

def main():
    # 输入和输出路径
    md_path = r"C:\Users\mingh\.openclaw\workspace\skills\business-expert\scripts\reports_fixed\CBD_经营分析报告_20260317.md"
    docx_path = r"C:\Users\mingh\.openclaw\workspace\skills\business-expert\scripts\reports_fixed\CBD_经营分析报告_20260317.docx"
    
    print(f"Markdown文件: {md_path}")
    print(f"Word输出文件: {docx_path}")
    
    # 读取Markdown内容
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        print(f"成功读取Markdown文件，大小: {len(markdown_content)} 字符")
    except Exception as e:
        print(f"读取Markdown文件失败: {e}")
        return
    
    # 转换文档
    try:
        # 尝试使用markdown-to-word-skill
        converter = MarkdownToDocxConverter(debug=True)
        converter.convert(markdown_content)
        converter.doc.save(docx_path)
        print(f"✅ 使用markdown-to-word-skill转换成功: {docx_path}")
    except Exception as e:
        print(f"使用markdown-to-word-skill转换失败: {e}")
        print("尝试使用简单转换...")
        
        try:
            # 使用简单转换
            converter = SimpleMarkdownToDocxConverter()
            converter.convert(markdown_content, docx_path)
            print(f"✅ 使用简单转换成功: {docx_path}")
        except Exception as e2:
            print(f"简单转换也失败: {e2}")
            return
    
    # 检查文件是否生成
    if os.path.exists(docx_path):
        file_size = os.path.getsize(docx_path)
        print(f"✅ Word文档生成成功!")
        print(f"   文件大小: {file_size:,} 字节")
        print(f"   文件路径: {docx_path}")
    else:
        print("❌ Word文档生成失败")

if __name__ == "__main__":
    main()